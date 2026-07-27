"""Parser tests (T-203, R-34).

Every fixture is generated in-process by the same libraries that parse it — PyMuPDF
writes the PDFs, python-docx the DOCXs, `zipfile` the bombs — so the suite carries no
binary files and stays runnable on a clean checkout.

The caps are exercised by shrinking :class:`ParserSettings` rather than by building
enormous inputs: a 5,001-page PDF proves nothing a 3-page PDF with ``max_pages=2``
does not, and it would add minutes to every run.
"""

from __future__ import annotations

import io
import zipfile
from pathlib import Path

import pymupdf
import pytest
from docx import Document as new_docx

from app.config import ParserSettings
from app.ingestion.parsers import (
    PREPROCESSING_VERSION,
    CorruptDocumentError,
    DocumentTooComplexError,
    EncryptedDocumentError,
    LocatorKind,
    NoExtractableTextError,
    UnsupportedDocumentError,
    parse_document,
    parse_document_sync,
)
from app.ingestion.parsers.base import split_text
from app.ingestion.parsers.text import is_blank, normalize

# --- fixtures -----------------------------------------------------------------


def _limits(**overrides) -> ParserSettings:
    return ParserSettings(**overrides)


def make_pdf(pages: int = 3, *, text: bool = True, encrypt: bool = False) -> bytes:
    document = pymupdf.open()
    for index in range(pages):
        page = document.new_page()
        if text:
            page.insert_text((72, 100), f"Page {index + 1} content")
    if encrypt:
        data = document.tobytes(
            encryption=pymupdf.PDF_ENCRYPT_AES_256, owner_pw="owner", user_pw="user"
        )
    else:
        data = document.tobytes()
    document.close()
    return data


def make_docx(*, headings: bool = True, table: bool = True, empty: bool = False) -> bytes:
    document = new_docx()
    if not empty:
        if headings:
            document.add_heading("Introduction", level=1)
        document.add_paragraph("Corpus ingests documents.")
        if headings:
            document.add_heading("Setup", level=2)
        document.add_paragraph("Install the thing.")
        if table:
            grid = document.add_table(rows=2, cols=3)
            values = [["Region", "Q3", "Q4"], ["EU", "12", "15"]]
            for row_index, row in enumerate(values):
                for col_index, value in enumerate(row):
                    grid.cell(row_index, col_index).text = value
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def docx_with_member(name: str, data: bytes) -> bytes:
    """A real DOCX with one extra member — the shape of a bomb hidden in a document."""
    original = io.BytesIO(make_docx())
    output = io.BytesIO()
    with (
        zipfile.ZipFile(original) as source,
        zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as target,
    ):
        for info in source.infolist():
            target.writestr(info, source.read(info.filename))
        target.writestr(name, data)
    return output.getvalue()


CSV_BYTES = b'company,region,revenue\nAcme,EU,12\nInitech,US,"line1\nline2"\nGlobex,APAC,7\n'

MD_BYTES = b"""# Guide

Intro with **bold** and a [label](https://example.com/secret).

## Setup

- first item
- second item

```bash
uv run pytest
```

| Region | Q3 |
| ------ | -- |
| EU     | 12 |

<script>alert(1)</script>

Trailing paragraph.
"""


# --- normalisation ------------------------------------------------------------


def test_normalize_is_idempotent() -> None:
    messy = "Title\r\n\r\n\r\n\r\nBody\ttext   \n\n\n\nEnd"
    once = normalize(messy)
    assert normalize(once) == once


def test_normalize_line_endings_and_blank_runs() -> None:
    assert normalize("a\r\nb\r\rc") == "a\nb\n\nc"
    assert normalize("a\n\n\n\n\nb") == "a\n\nb"


def test_normalize_strips_controls_and_folds_odd_spaces() -> None:
    assert normalize("a\x00\x07b") == "ab"
    assert normalize("a b　c") == "a b c"
    assert normalize("wo​rd­here") == "wordhere"


def test_normalize_trailing_whitespace_and_empty() -> None:
    assert normalize("line   \nnext\t\n") == "line\nnext"
    assert normalize("") == ""
    assert is_blank("   \n ")


def test_normalize_keeps_nfc_equivalent_forms_identical() -> None:
    # "é" precomposed vs. e + combining acute must hash the same.
    assert normalize("café") == normalize("café")


# --- split_text ---------------------------------------------------------------


def test_split_text_respects_the_ceiling() -> None:
    text = "\n".join(f"line {index}" for index in range(200))
    parts = split_text(text, 100)
    assert len(parts) > 1
    assert all(len(part) <= 100 for part in parts)
    assert "".join(part.replace("\n", "") for part in parts) == text.replace("\n", "")


def test_split_text_hard_slices_an_overlong_line() -> None:
    parts = split_text("x" * 250, 100)
    assert [len(part) for part in parts] == [100, 100, 50]


def test_split_text_returns_short_text_untouched() -> None:
    assert split_text("short", 100) == ["short"]


# --- PDF ----------------------------------------------------------------------


def test_pdf_blocks_are_pages_with_1_based_locators() -> None:
    parsed = parse_document_sync(make_pdf(3), filename="report.pdf")

    assert parsed.suffix == ".pdf"
    assert parsed.page_count == 3
    assert len(parsed.blocks) == 3
    assert [block.locator.page for block in parsed.blocks] == [1, 2, 3]
    assert [block.locator.label for block in parsed.blocks] == ["p. 1", "p. 2", "p. 3"]
    assert all(block.locator.kind is LocatorKind.PAGE for block in parsed.blocks)
    assert parsed.blocks[1].text == "Page 2 content"
    assert parsed.preprocessing_version == PREPROCESSING_VERSION


def test_pdf_locator_metadata_shape() -> None:
    parsed = parse_document_sync(make_pdf(1), filename="report.pdf")
    assert parsed.blocks[0].locator.as_metadata() == {"kind": "page", "label": "p. 1", "page": 1}


def test_pdf_without_text_fails_rather_than_ingesting_empty() -> None:
    with pytest.raises(NoExtractableTextError, match="OCR"):
        parse_document_sync(make_pdf(2, text=False), filename="scan.pdf")


def test_pdf_password_protected_is_rejected() -> None:
    with pytest.raises(EncryptedDocumentError):
        parse_document_sync(make_pdf(1, encrypt=True), filename="secret.pdf")


def test_pdf_with_a_valid_header_but_junk_body_is_corrupt() -> None:
    """Exactly what gets past the upload sniffer: `%PDF-` at offset 0, nothing behind it."""
    with pytest.raises(CorruptDocumentError, match="could not be opened"):
        parse_document_sync(b"%PDF-1.7\n" + b"\x01\x02garbage" * 20, filename="broken.pdf")


def test_pdf_damaged_but_openable_fails_as_empty_not_as_a_crash() -> None:
    """PyMuPDF recovers aggressively: a lightly truncated file opens with zero pages.

    Either terminal code is honest here; what must not happen is a raw MuPDF error
    escaping the parser, which would surface as a worker crash rather than a document
    the user can see has failed.
    """
    with pytest.raises(NoExtractableTextError):
        parse_document_sync(make_pdf(1)[:120], filename="clipped.pdf")


def test_pdf_page_cap() -> None:
    with pytest.raises(DocumentTooComplexError, match="pages"):
        parse_document_sync(make_pdf(3), filename="report.pdf", limits=_limits(max_pages=2))


def test_pdf_character_budget() -> None:
    with pytest.raises(DocumentTooComplexError, match="characters"):
        parse_document_sync(
            make_pdf(3), filename="report.pdf", limits=_limits(max_extracted_chars=20)
        )


# --- DOCX ---------------------------------------------------------------------


def test_docx_sections_carry_nested_heading_paths() -> None:
    parsed = parse_document_sync(make_docx(), filename="handbook.docx")

    assert parsed.suffix == ".docx"
    assert parsed.page_count is None  # R-34: DOCX stores no pagination
    paths = [block.locator.section_path for block in parsed.blocks]
    assert paths == [("Introduction",), ("Introduction", "Setup")]
    assert parsed.blocks[1].locator.label == "§ Introduction › Setup"
    assert all(block.locator.kind is LocatorKind.SECTION for block in parsed.blocks)


def test_docx_tables_survive_and_keep_row_adjacency() -> None:
    parsed = parse_document_sync(make_docx(), filename="handbook.docx")
    text = parsed.text
    assert "Region | Q3 | Q4" in text
    assert "EU | 12 | 15" in text


def test_docx_table_stays_in_body_order_after_its_paragraph() -> None:
    parsed = parse_document_sync(make_docx(), filename="handbook.docx")
    body = parsed.blocks[-1].text
    assert body.index("Install the thing.") < body.index("Region | Q3 | Q4")


def test_docx_without_headings_falls_back_to_a_paragraph_label() -> None:
    parsed = parse_document_sync(make_docx(headings=False), filename="notes.docx")
    assert len(parsed.blocks) == 1
    assert parsed.blocks[0].locator.section_path == ()
    assert parsed.blocks[0].locator.label.startswith("paragraph ")


def test_deep_heading_paths_shorten_the_label_but_not_the_metadata() -> None:
    document = new_docx()
    for level, title in enumerate(("One", "Two", "Three", "Four"), start=1):
        document.add_heading(title, level=level)
    document.add_paragraph("leaf text")
    buffer = io.BytesIO()
    document.save(buffer)

    locator = parse_document_sync(buffer.getvalue(), filename="deep.docx").blocks[-1].locator
    assert locator.section_path == ("One", "Two", "Three", "Four")
    assert locator.label == "§ Three › Four"
    assert locator.as_metadata()["section_path"] == ["One", "Two", "Three", "Four"]


def test_docx_heading_level_jump_does_not_pad_the_path() -> None:
    document = new_docx()
    document.add_heading("Top", level=1)
    document.add_paragraph("a")
    document.add_heading("Deep", level=3)
    document.add_paragraph("b")
    buffer = io.BytesIO()
    document.save(buffer)

    parsed = parse_document_sync(buffer.getvalue(), filename="jump.docx")
    assert parsed.blocks[-1].locator.section_path == ("Top", "Deep")


def test_docx_empty_document_fails() -> None:
    with pytest.raises(NoExtractableTextError):
        parse_document_sync(make_docx(empty=True), filename="blank.docx")


def test_docx_not_a_zip_is_corrupt() -> None:
    with pytest.raises(CorruptDocumentError, match="ZIP"):
        parse_document_sync(b"PK\x03\x04not really a zip", filename="broken.docx")


def test_docx_expanded_size_cap() -> None:
    with pytest.raises(DocumentTooComplexError, match="expands to"):
        parse_document_sync(
            make_docx(), filename="handbook.docx", limits=_limits(docx_max_expanded_bytes=500)
        )


def test_docx_member_count_cap() -> None:
    with pytest.raises(DocumentTooComplexError, match="parts"):
        parse_document_sync(
            make_docx(), filename="handbook.docx", limits=_limits(docx_max_members=2)
        )


def test_docx_compression_ratio_cap_catches_a_zip_bomb() -> None:
    # 8 MB of zeros deflates to a few kilobytes — a ratio no real document reaches.
    bomb = docx_with_member("word/media/payload.bin", b"\0" * (8 * 1024 * 1024))
    with pytest.raises(DocumentTooComplexError, match="ratio"):
        parse_document_sync(bomb, filename="bomb.docx")


def test_docx_bomb_is_rejected_before_decompression() -> None:
    """The cap must trip on the central directory, not after expanding the member."""
    bomb = docx_with_member("word/media/payload.bin", b"\0" * (8 * 1024 * 1024))
    assert len(bomb) < 1024 * 1024  # the file on disk is small; only its expansion is not
    with pytest.raises(DocumentTooComplexError):
        parse_document_sync(bomb, filename="bomb.docx")


# --- CSV ----------------------------------------------------------------------


def test_csv_header_is_detected_and_repeated_in_every_block() -> None:
    parsed = parse_document_sync(
        CSV_BYTES, filename="sales.csv", limits=_limits(csv_rows_per_block=2)
    )

    assert parsed.suffix == ".csv"
    assert parsed.page_count is None
    assert len(parsed.blocks) == 2
    assert all(block.text.startswith("company | region | revenue") for block in parsed.blocks)


def test_csv_row_ranges_are_contiguous_and_header_exclusive() -> None:
    parsed = parse_document_sync(
        CSV_BYTES, filename="sales.csv", limits=_limits(csv_rows_per_block=2)
    )

    ranges = [(block.locator.row_start, block.locator.row_end) for block in parsed.blocks]
    assert ranges == [(1, 2), (3, 3)]
    assert parsed.blocks[0].locator.kind is LocatorKind.ROWS
    assert parsed.blocks[0].locator.label == "rows 1–2"
    assert parsed.blocks[1].locator.label == "row 3"


def test_csv_quoted_newline_is_one_record_not_two() -> None:
    parsed = parse_document_sync(CSV_BYTES, filename="sales.csv")
    assert parsed.blocks[0].locator.row_end == 3  # three data rows, not four
    assert "line1 line2" in parsed.blocks[0].text


def test_csv_first_row_of_numbers_is_data_not_a_header() -> None:
    parsed = parse_document_sync(b"1,2,3\n4,5,6\n", filename="matrix.csv")
    assert parsed.blocks[0].locator.row_end == 2
    assert parsed.blocks[0].text.startswith("1 | 2 | 3")


def test_csv_repeated_labels_are_not_a_header() -> None:
    parsed = parse_document_sync(b"a,a,b\nx,y,z\n", filename="dupes.csv")
    assert parsed.blocks[0].locator.row_end == 2


@pytest.mark.parametrize(
    "payload",
    [
        b"company;region\nAcme;EU\n",  # semicolon
        b"company\tregion\nAcme\tEU\n",  # tab
        b"company|region\nAcme|EU\n",  # pipe
    ],
)
def test_csv_dialects_are_sniffed(payload: bytes) -> None:
    parsed = parse_document_sync(payload, filename="data.csv")
    assert "company | region" in parsed.blocks[0].text
    assert "Acme | EU" in parsed.blocks[0].text


def test_csv_row_cap() -> None:
    payload = b"a,b\n" + b"1,2\n" * 20
    with pytest.raises(DocumentTooComplexError, match="rows"):
        parse_document_sync(payload, filename="big.csv", limits=_limits(csv_max_rows=5))


def test_csv_column_cap() -> None:
    payload = ",".join(f"c{index}" for index in range(20)).encode() + b"\n"
    with pytest.raises(DocumentTooComplexError, match="columns"):
        parse_document_sync(payload, filename="wide.csv", limits=_limits(csv_max_columns=5))


def test_csv_oversized_field_is_a_cap_breach_not_a_corruption() -> None:
    payload = b"a,b\n" + b"x" * 500 + b",2\n"
    with pytest.raises(DocumentTooComplexError, match="oversized field"):
        parse_document_sync(payload, filename="fat.csv", limits=_limits(max_block_chars=100))


def test_csv_field_size_limit_is_restored_after_parsing() -> None:
    import csv as stdlib_csv

    before = stdlib_csv.field_size_limit()
    parse_document_sync(CSV_BYTES, filename="sales.csv", limits=_limits(max_block_chars=4096))
    assert stdlib_csv.field_size_limit() == before


def test_csv_non_utf8_is_corrupt() -> None:
    with pytest.raises(CorruptDocumentError, match="UTF-8"):
        parse_document_sync(b"a,b\n\xff\xfe,2\n", filename="latin.csv")


def test_csv_header_only_has_no_data_rows() -> None:
    with pytest.raises(NoExtractableTextError):
        parse_document_sync(b"company,region\n", filename="empty.csv")


# --- Markdown -----------------------------------------------------------------


def test_markdown_sections_and_line_ranges() -> None:
    parsed = parse_document_sync(MD_BYTES, filename="guide.md")

    assert parsed.suffix == ".md"
    assert parsed.page_count is None
    assert [block.locator.section_path for block in parsed.blocks] == [
        ("Guide",),
        ("Guide", "Setup"),
    ]
    first, second = parsed.blocks
    assert first.locator.line_start == 1
    assert first.locator.line_end < second.locator.line_start
    assert first.locator.as_metadata()["line_start"] == 1


def test_markdown_strips_syntax_but_keeps_content() -> None:
    parsed = parse_document_sync(MD_BYTES, filename="guide.md")
    text = parsed.text
    assert "Intro with bold and a label." in text
    assert "**" not in text


def test_markdown_drops_link_urls() -> None:
    parsed = parse_document_sync(MD_BYTES, filename="guide.md")
    assert "example.com" not in parsed.text


def test_markdown_refuses_raw_html() -> None:
    parsed = parse_document_sync(MD_BYTES, filename="guide.md")
    assert "script" not in parsed.text.lower()


def test_markdown_keeps_code_fences_and_groups_lists_and_tables() -> None:
    parsed = parse_document_sync(MD_BYTES, filename="guide.md")
    body = parsed.blocks[-1].text
    assert "uv run pytest" in body
    assert "- first item\n- second item" in body  # one block, not two paragraphs
    assert "Region | Q3\nEU | 12" in body


def test_markdown_without_headings_uses_a_line_range_label() -> None:
    parsed = parse_document_sync(b"just a paragraph\n\nand another\n", filename="plain.md")
    assert len(parsed.blocks) == 1
    assert parsed.blocks[0].locator.section_path == ()
    assert parsed.blocks[0].locator.label.startswith("lines ")


def test_markdown_empty_fails() -> None:
    with pytest.raises(NoExtractableTextError):
        parse_document_sync(b"<!-- nothing but a comment -->\n", filename="empty.md")


def test_markdown_non_utf8_is_corrupt() -> None:
    with pytest.raises(CorruptDocumentError, match="UTF-8"):
        parse_document_sync(b"# heading\n\xff\xfe\n", filename="bad.md")


# --- dispatch -----------------------------------------------------------------


@pytest.mark.parametrize(
    ("payload_factory", "filename", "suffix"),
    [
        (lambda: make_pdf(1), "a.pdf", ".pdf"),
        (make_docx, "a.docx", ".docx"),
        (lambda: CSV_BYTES, "a.csv", ".csv"),
        (lambda: MD_BYTES, "a.md", ".md"),
    ],
)
def test_dispatch_by_extension(payload_factory, filename: str, suffix: str) -> None:
    parsed = parse_document_sync(payload_factory(), filename=filename)
    assert parsed.suffix == suffix
    assert parsed.blocks
    assert parsed.char_count == sum(len(block.text) for block in parsed.blocks)
    assert [block.order for block in parsed.blocks] == list(range(len(parsed.blocks)))


def test_dispatch_is_case_insensitive() -> None:
    assert parse_document_sync(CSV_BYTES, filename="SALES.CSV").suffix == ".csv"


def test_unknown_extension_is_rejected() -> None:
    with pytest.raises(UnsupportedDocumentError):
        parse_document_sync(b"anything", filename="notes.txt")


def test_empty_payload_is_rejected() -> None:
    with pytest.raises(CorruptDocumentError, match="empty"):
        parse_document_sync(b"", filename="notes.csv")


async def test_async_facade_matches_sync() -> None:
    payload = make_pdf(2)
    asynchronous = await parse_document(payload, filename="report.pdf")
    synchronous = parse_document_sync(payload, filename="report.pdf")
    assert asynchronous == synchronous


def test_blocks_never_exceed_the_block_ceiling() -> None:
    limits = _limits(max_block_chars=40)
    for payload, filename in (
        (make_pdf(2), "a.pdf"),
        (make_docx(), "a.docx"),
        (CSV_BYTES, "a.csv"),
        (MD_BYTES, "a.md"),
    ):
        parsed = parse_document_sync(payload, filename=filename, limits=limits)
        assert all(len(block.text) <= 40 for block in parsed.blocks), filename


def test_split_blocks_keep_their_parent_locator() -> None:
    parsed = parse_document_sync(make_pdf(1), filename="a.pdf", limits=_limits(max_block_chars=5))
    assert len(parsed.blocks) > 1
    assert {block.locator.page for block in parsed.blocks} == {1}


# --- placement (R-31: parse in the worker, never in the API process) ----------


def test_no_api_module_imports_the_parsers() -> None:
    api_dir = Path(__file__).resolve().parents[1] / "app" / "api"
    offenders = [
        path.name
        for path in api_dir.glob("*.py")
        if "ingestion.parsers" in path.read_text(encoding="utf-8")
    ]
    assert not offenders, (
        f"{offenders} import the parsers into the API process — R-31 (§8.12) requires "
        "parsing to happen in the ingestion worker"
    )
