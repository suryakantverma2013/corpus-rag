"""Tabular structure (T-219, R-88(5)/(6) §8.78; T-223, R-89 — FR-ING-08).

**Two mechanisms, one output contract, and the split is the thing to keep straight.** A PDF
table is *detected*: PyMuPDF looks for ruling lines, so detection is heuristic, it is floored
against false positives by `PARSER_TABLE_*`, and everything it rejects stays readable as
ordinary page text. A DOCX or Markdown table is *declared*: the format says where the table is
and which cells are in which column, so there is nothing to guess and nothing to floor against
— what the same two floors happen to catch there is the layout table (R-89). Both produce the
same thing: one `ParsedBlock` marked `table`, on its container's own locator, whose first line
is its header verbatim when it has one.

Fixtures are generated in-process by the same library that parses them, per T-203's rule: the
suite carries no binary files and stays runnable on a clean checkout. A "table" here is a real
ruled grid, because the shipped detection strategy is `lines_strict` — a whitespace-aligned
fixture would pass or fail for reasons unrelated to the code under test.

**The first assertion in this module is the one the board line asks for first**: a cell's text
must not reach the index twice. It is the failure that is invisible downstream — both chunks are
well-formed, both embed, and retrieval simply answers one passage under two citations.

Absence assertions carry real weight here. "The table's region left no text behind" is the only
way to state the exclusion, and a test that checks only the table block passes just as happily
against a parser that emits the table *and* leaves every cell in the page text.
"""

from __future__ import annotations

import io
import subprocess
import sys
import textwrap
from pathlib import Path

import pymupdf
import pytest
from docx import Document as new_docx
from docx.oxml.ns import qn

from app.config import ParserSettings, Settings
from app.ingestion.parsers import parse_document_sync
from app.ingestion.parsers import pdf as pdf_parser
from app.ingestion.parsers.base import (
    DocumentTooComplexError,
    Extraction,
    LocatorKind,
    ParsedDocument,
)
from app.ingestion.parsers.tables import DetectedTable, compose_page, detect_tables
from app.ingestion.parsers.text import normalize
from app.services.ocr import OcrClient, OcrPageResult, OcrWord

PROSE_ABOVE = "Quarterly results follow."
PROSE_BELOW = "Notes appear after the table."
HEADER = ["Region", "Q3", "Q4"]
BODY = [["EU", "12", "15"], ["US", "20", "25"]]
GRID = [HEADER, *BODY]


# --- fixtures -----------------------------------------------------------------


def _limits(**overrides) -> ParserSettings:
    """Parser limits with recognition **pinned off**, never inherited from the environment.

    FR-ING-08 runs on the text layer and needs no sidecar (R-88(5)); leaving `ocr_enabled` to
    `ParserSettings()` would let a developer's `.env` arm it inside this module and make the
    determinism tests depend on a container. T-218 learned this the same way.
    """
    return ParserSettings(**{"ocr_enabled": False, **overrides})


def make_pdf_with_tables(
    grids: list[list[list[str]]] | None = None,
    *,
    external_header: bool = False,
    prose: bool = True,
    ruled: bool = True,
    page_height: float = 792.0,
    cell_height: float = 20.0,
    gap: float = 30.0,
) -> bytes:
    """A born-digital page carrying ruled grids, optionally with prose above and below.

    ``external_header`` draws each grid's first row as unruled text *above* the box, which is
    what makes `find_tables` report `header.external` — a real layout, and the one where the
    header sits outside `table.bbox`.
    """
    grids = [GRID] if grids is None else grids
    document = pymupdf.open()
    page = document.new_page(width=595.0, height=page_height)
    if prose:
        page.insert_text((72, 80), PROSE_ABOVE)

    top = 120.0
    width = 120.0
    for grid in grids:
        rows = grid[1:] if external_header else grid
        if external_header:
            for column, name in enumerate(grid[0]):
                page.insert_text((72 + column * width + 3, top - 6), name)
        for row_index, row in enumerate(rows):
            for column, value in enumerate(row):
                rect = pymupdf.Rect(
                    72 + column * width,
                    top + row_index * cell_height,
                    72 + (column + 1) * width,
                    top + (row_index + 1) * cell_height,
                )
                if ruled:
                    page.draw_rect(rect)
                page.insert_text((rect.x0 + 3, rect.y1 - 6), value)
        top += len(rows) * cell_height + gap

    if prose:
        page.insert_text((72, top + 20), PROSE_BELOW)
    data = document.tobytes()
    document.close()
    return data


def make_scanned_table_pdf(dpi: int = 150) -> bytes:
    """The same ruled table, flattened to a raster: no text layer at all."""
    rendered = pymupdf.open(stream=make_pdf_with_tables(), filetype="pdf")
    out = pymupdf.open()
    raster = rendered[0].get_pixmap(dpi=dpi).tobytes("png")
    page = out.new_page()
    page.insert_image(page.rect, stream=raster)
    rendered.close()
    data = out.tobytes()
    out.close()
    return data


class _FakeOcr(OcrClient):
    """A recogniser that replies with reading-order text. Subclass, never a Protocol (R-88(1))."""

    def __init__(self, text: str) -> None:
        super().__init__(Settings())
        self._text = text
        self.calls = 0

    def recognize(self, image: bytes) -> OcrPageResult:  # noqa: ARG002 — the raster is not read
        self.calls += 1
        return OcrPageResult(
            text=self._text,
            words=tuple(OcrWord(text=word, confidence=95.0) for word in self._text.split()),
            engine_version="5.5.0",
            languages=("eng",),
        )


def _parse(payload: bytes, *, limits: ParserSettings | None = None, **kwargs) -> ParsedDocument:
    return pdf_parser.parse(payload, limits=limits or _limits(), **kwargs)


def _blocks_of(parsed: ParsedDocument, extraction: Extraction) -> list:
    return [block for block in parsed.blocks if block.extraction is extraction]


def _text_of(parsed: ParsedDocument, extraction: Extraction) -> str:
    return "\n".join(block.text for block in _blocks_of(parsed, extraction))


def _first_page(payload: bytes):
    document = pymupdf.open(stream=payload, filetype="pdf")
    return document, document.load_page(0)


# --- the failure this feature exists to avoid ---------------------------------


def test_a_tables_cells_are_not_also_indexed_as_page_text() -> None:
    """Indexed twice, one passage answers under two citations and both chunks look healthy."""
    parsed = _parse(make_pdf_with_tables())

    page_text = _text_of(parsed, Extraction.TEXT)
    for cell in ("Region", "Q3", "EU", "12", "US", "25"):
        assert cell not in page_text, f"{cell!r} survived in the page text as well as the table"
    assert PROSE_ABOVE in page_text
    assert PROSE_BELOW in page_text


def test_an_external_header_is_not_left_behind_in_the_page_text() -> None:
    """`find_tables` places an external header *outside* `table.bbox` — hence the union.

    Excluding the bounding box alone leaves the column names in the prose while the table block
    also carries them: the same double index, one row at a time, and only on the layouts where a
    header happens to sit above the ruling.
    """
    payload = make_pdf_with_tables(external_header=True)
    document, page = _first_page(payload)
    try:
        detected = detect_tables(page, limits=_limits())
        assert len(detected) == 1
        assert page.find_tables().tables[0].header.external is True
    finally:
        document.close()

    parsed = _parse(payload)
    assert "Region" not in _text_of(parsed, Extraction.TEXT)
    assert _blocks_of(parsed, Extraction.TABLE)[0].text.startswith("Region | Q3 | Q4")


def test_an_internal_header_is_not_repeated_inside_its_own_block() -> None:
    """When the header is not external, `extract()[0]` **is** the header row."""
    parsed = _parse(make_pdf_with_tables())
    text = _blocks_of(parsed, Extraction.TABLE)[0].text
    assert text.count("Region | Q3 | Q4") == 1
    assert text.splitlines() == ["Region | Q3 | Q4", "EU | 12 | 15", "US | 20 | 25"]


# --- shape of what is emitted -------------------------------------------------


def test_a_table_is_one_block_on_its_pages_own_locator() -> None:
    """R-88(6): no fourth `LocatorKind`, and the header is declared rather than inferred."""
    parsed = _parse(make_pdf_with_tables())
    tables = _blocks_of(parsed, Extraction.TABLE)
    assert len(tables) == 1
    assert tables[0].locator.kind is LocatorKind.PAGE
    assert tables[0].locator.page == 1
    assert tables[0].header == "Region | Q3 | Q4"
    assert all(block.locator.kind is LocatorKind.PAGE for block in parsed.blocks)


def test_prose_above_and_below_a_table_keep_reading_order() -> None:
    """Interleaved, because those two paragraphs were never adjacent."""
    parsed = _parse(make_pdf_with_tables())
    assert [block.extraction for block in parsed.blocks] == [
        Extraction.TEXT,
        Extraction.TABLE,
        Extraction.TEXT,
    ]
    assert parsed.blocks[0].text == PROSE_ABOVE
    assert parsed.blocks[2].text == PROSE_BELOW
    assert [block.order for block in parsed.blocks] == [0, 1, 2]


def test_a_table_with_no_prose_still_yields_only_the_table() -> None:
    parsed = _parse(make_pdf_with_tables(prose=False))
    assert [block.extraction for block in parsed.blocks] == [Extraction.TABLE]


# --- a page without tables is untouched ---------------------------------------


def test_a_page_with_no_table_extracts_exactly_as_before() -> None:
    """Every PDF in the corpus goes through this path; it must be byte-identical to today's."""
    payload = make_pdf_with_tables(ruled=False)
    document, page = _first_page(payload)
    try:
        assert detect_tables(page, limits=_limits()) == []
        expected = normalize(page.get_text("text", sort=True))
    finally:
        document.close()

    parsed = _parse(payload)
    assert [block.extraction for block in parsed.blocks] == [Extraction.TEXT]
    assert parsed.blocks[0].text == expected


def test_an_unruled_table_is_not_detected_and_is_the_recorded_limitation() -> None:
    """The pinned `lines_strict` strategy needs ruling lines. Recorded, not discovered later."""
    parsed = _parse(make_pdf_with_tables(ruled=False))
    assert _blocks_of(parsed, Extraction.TABLE) == []
    assert "Region" in _text_of(parsed, Extraction.TEXT)


# --- thresholds degrade rather than mangle ------------------------------------


def test_a_table_below_the_row_floor_is_left_as_ordinary_text() -> None:
    parsed = _parse(make_pdf_with_tables(), limits=_limits(table_min_rows=4))
    assert _blocks_of(parsed, Extraction.TABLE) == []
    assert "Region" in _text_of(parsed, Extraction.TEXT)
    assert "12" in _text_of(parsed, Extraction.TEXT)


def test_a_table_below_the_column_floor_is_left_as_ordinary_text() -> None:
    parsed = _parse(make_pdf_with_tables(), limits=_limits(table_min_columns=4))
    assert _blocks_of(parsed, Extraction.TABLE) == []
    assert "Region" in _text_of(parsed, Extraction.TEXT)


def test_the_per_page_cap_leaves_the_surplus_tables_text_in_the_page() -> None:
    """A cap that also hid the surplus regions would delete their content instead."""
    payload = make_pdf_with_tables([GRID, GRID, GRID], cell_height=14.0, gap=20.0)
    parsed = _parse(payload, limits=_limits(table_max_per_page=1))
    assert len(_blocks_of(parsed, Extraction.TABLE)) == 1
    # The two tables that did not make the cap are still readable as text.
    assert _text_of(parsed, Extraction.TEXT).count("Region") == 2


def test_the_cap_keeps_the_tables_nearest_the_top_of_the_page() -> None:
    payload = make_pdf_with_tables([GRID, GRID, GRID], cell_height=14.0, gap=20.0)
    document, page = _first_page(payload)
    try:
        detected = detect_tables(page, limits=_limits(table_max_per_page=2))
        tops = [table.top for table in detected]
    finally:
        document.close()
    assert len(detected) == 2
    assert tops == sorted(tops)


# --- failing open (R-88(9)'s shape, one requirement over) ----------------------


def test_a_detection_fault_never_fails_a_page_that_extracted_perfectly(monkeypatch) -> None:
    """Layout analysis is an enrichment over text the page has already yielded.

    `find_tables` is a large heuristic over the most hostile input this system accepts, and it
    sits *outside* the page-level handler that makes an extraction fault a clean
    `CORRUPT_DOCUMENT`. Without its own handler a fault here escapes as an unclassified crash
    and fails a document whose characters came out perfectly — the asymmetry R-88(9) rejects
    for recognition, arriving one requirement later by a different door.
    """

    def explode(*args, **kwargs):
        raise RuntimeError("layout analysis fell over")

    monkeypatch.setattr(pdf_parser, "detect_tables", explode)
    parsed = _parse(make_pdf_with_tables())

    assert _blocks_of(parsed, Extraction.TABLE) == []
    text = _text_of(parsed, Extraction.TEXT)
    assert PROSE_ABOVE in text
    assert "Region" in text, "the page degraded to no text at all, not to plain text"


def test_a_budget_breach_inside_a_table_page_is_still_terminal(monkeypatch) -> None:
    """Two claims at once, because one parse proves both.

    A table's characters are charged against `PARSER_MAX_EXTRACTED_CHARS` like any others —
    without that a table-dense document would bypass the ceiling entirely. And the emission sits
    **outside** the fail-open handler on purpose: `CharBudget` is R-34's caps-reject-never-truncate
    rule, and a handler written for detection faults must not swallow a terminal one.
    """
    from app.ingestion.parsers.base import DocumentTooComplexError

    with pytest.raises(DocumentTooComplexError):
        _parse(make_pdf_with_tables(), limits=_limits(max_extracted_chars=30))


# --- R-88(5): the scanned-table gap, pinned as known --------------------------


def test_a_scanned_table_yields_recognised_text_and_no_grid() -> None:
    """OCR recovers characters; it does not recover a structure the page never encoded.

    This is the case a user tests early, so R-88(5) put it in the ruling rather than a docstring
    and this pins it: a scanned table ingests as reading-order recognised text, marked `ocr`, with
    no table block anywhere.
    """
    fake = _FakeOcr("Region Q3 Q4 EU 12 15 US 20 25")
    parsed = _parse(make_scanned_table_pdf(), limits=_limits(ocr_enabled=True), recognizer=fake)
    assert fake.calls == 1
    assert _blocks_of(parsed, Extraction.TABLE) == []
    assert _blocks_of(parsed, Extraction.OCR)
    assert "Region Q3 Q4" in _text_of(parsed, Extraction.OCR)


def test_a_scanned_page_is_never_table_detected() -> None:
    """Detection reads the text layer, so a page without one must not even be asked."""
    document, page = _first_page(make_scanned_table_pdf())
    try:
        assert page.get_text("text").strip() == ""
        assert detect_tables(page, limits=_limits()) == []
    finally:
        document.close()


# --- determinism (R-88(1)) ----------------------------------------------------


def test_two_parses_of_one_file_are_byte_identical() -> None:
    """Emitted text feeds `embedding_fingerprint` through `chunk_text`; anything that varies
    between two passes over one file leaves vector reuse permanently empty (R-88(1))."""
    payload = make_pdf_with_tables([GRID, GRID], cell_height=14.0, gap=20.0)
    assert _parse(payload) == _parse(payload)


def test_composed_items_are_ordered_by_geometry() -> None:
    """`compose_page` is what decides block order, so it is what this asserts.

    `detect_tables` sorts too, but PyMuPDF already hands back geometric order — measured, drawing
    the lower grid into the content stream first — so a test aimed there would pass with the sort
    deleted and certify nothing. What the sort in `detect_tables` really pins is the cap's meaning
    ("the topmost N"), which `test_the_cap_keeps_the_tables_nearest_the_top_of_the_page` covers.
    """
    document, page = _first_page(make_pdf_with_tables([GRID, GRID], cell_height=14.0, gap=20.0))
    try:
        composed = compose_page(page, detect_tables(page, limits=_limits()))
    finally:
        document.close()

    table_tops = [item.rect.y0 for item in composed if isinstance(item, DetectedTable)]
    assert len(table_tops) == 2
    assert table_tops == sorted(table_tops)

    # **The load-bearing pair.** Asserting only "the tables are in order" passes against a
    # `compose_page` that appends every table after every prose run — the prose is already sorted
    # on the way in and `detect_tables` hands the tables over in order, so both halves hold by
    # accident. What the sort actually decides is where prose that *follows* a table lands: with
    # it, `PROSE_BELOW` is the last item; without it, it is folded into the opening run and the
    # two paragraphs either side of the tables become one.
    assert isinstance(composed[0], str) and PROSE_ABOVE in composed[0]
    assert PROSE_BELOW not in composed[0]
    assert isinstance(composed[-1], str) and PROSE_BELOW in composed[-1]


def test_the_layout_advisory_never_reaches_stdout() -> None:
    """PyMuPDF `print()`s a suggestion on first use; the worker's log stream is JSON on stdout.

    One un-parseable English line in a structured log is a defect an operator meets long after the
    run that produced it, and `set_messages()` does not suppress this one — only
    `no_recommend_layout()` does.

    **In a subprocess, and that is the whole point.** PyMuPDF fires the advisory once per
    *process*, so an in-process assertion would pass whenever anything earlier in the session had
    already spent it: vacuous, and vacuous depending on test ordering. A fresh interpreter asks the
    question exactly once, of `tables.py`'s own import-time call, which is the line that can be
    deleted. (Re-arming the private flag and reloading the module was tried first and is worse: a
    reload rebinds `DetectedTable` inside the live module dict while `pdf.py` keeps the old alias,
    so every later test in the session sees `isinstance` fail.)
    """
    script = textwrap.dedent(
        """
        import pymupdf
        from app.config import ParserSettings
        from app.ingestion.parsers.tables import detect_tables

        document = pymupdf.open()
        page = document.new_page()
        for row in range(2):
            for column in range(2):
                page.draw_rect(pymupdf.Rect(72 + column * 60, 100 + row * 20,
                                            132 + column * 60, 120 + row * 20))
                page.insert_text((75 + column * 60, 114 + row * 20), "x")
        data = document.tobytes()
        document.close()

        opened = pymupdf.open(stream=data, filetype="pdf")
        detect_tables(opened.load_page(0), limits=ParserSettings(ocr_enabled=False))
        opened.close()
        """
    )
    result = subprocess.run(  # noqa: S603 — our own interpreter, our own script
        [sys.executable, "-c", script],
        capture_output=True,
        text=True,
        cwd=Path(__file__).resolve().parents[1],
        check=True,
    )
    assert result.stdout == "", f"PyMuPDF wrote to the worker's log stream: {result.stdout!r}"


# --- block-level splitting ----------------------------------------------------


def test_a_table_split_by_the_block_ceiling_repeats_its_header_on_every_part() -> None:
    """`split_text` may cut a huge table before the chunker ever sees it (R-88(6))."""
    rows = [HEADER, *[[f"R{index}", "12", "15"] for index in range(60)]]
    payload = make_pdf_with_tables([rows], cell_height=10.0, page_height=1400.0)
    parsed = _parse(payload, limits=_limits(max_block_chars=200))

    parts = _blocks_of(parsed, Extraction.TABLE)
    assert len(parts) > 1
    for part in parts:
        assert part.text.startswith("Region | Q3 | Q4")
        assert part.header == "Region | Q3 | Q4"
    # T-223: `startswith` passes just as happily against a part 0 carrying the header twice.
    assert parts[0].text.count("Region | Q3 | Q4") == 1


# --- settings -----------------------------------------------------------------


@pytest.mark.parametrize(
    "override",
    [
        {"table_min_rows": 0},
        {"table_min_columns": 0},
        {"table_max_per_page": 0},
    ],
)
def test_a_table_threshold_below_one_is_refused_at_boot(override) -> None:
    """Zero is not "unbounded" for any of these — it is a feature silently switched off."""
    with pytest.raises(ValueError):
        ParserSettings(**override)


# --- declared tables: DOCX and Markdown (T-223, R-89) -------------------------
#
# Everything above this line is *detection*. Everything below is *declaration*: the format
# states the grid, so these tests never build a ruled rectangle and never touch a threshold
# except to prove the floor is applied.


def _docx(
    *,
    rows: int = 3,
    cols: int = 3,
    header_flag: str | None = None,
    table_look: bool = True,
    blank_first_row: bool = False,
) -> bytes:
    """A DOCX whose § "Data" holds prose, one table, then prose."""
    document = new_docx()
    document.add_heading("Data", level=1)
    document.add_paragraph(PROSE_ABOVE)
    grid = document.add_table(rows=rows, cols=cols)
    for row_index in range(rows):
        for col_index in range(cols):
            if blank_first_row and row_index == 0:
                continue
            grid.cell(row_index, col_index).text = (
                HEADER[col_index] if row_index == 0 else f"R{row_index}C{col_index}"
            )

    properties = grid._tbl.tblPr
    look = properties.find(qn("w:tblLook"))
    if not table_look and look is not None:
        properties.remove(look)
    if header_flag is not None:
        row_properties = grid.rows[0]._tr.get_or_add_trPr()
        flag = row_properties.makeelement(qn("w:tblHeader"), {})
        if header_flag:
            flag.set(qn("w:val"), header_flag)
        row_properties.append(flag)

    document.add_paragraph(PROSE_BELOW)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


MD_TABLE = b"""# Data

Quarterly results follow.

| Region | Q3 | Q4 |
| ------ | -- | -- |
| EU     | 12 | 15 |
| US     | 20 | 25 |

Notes appear after the table.
"""


def _declared(payload: bytes, name: str, **overrides) -> ParsedDocument:
    return parse_document_sync(payload, filename=name, limits=_limits(**overrides))


# --- DOCX ---------------------------------------------------------------------


def test_a_docx_table_is_its_own_block_marked_table() -> None:
    parsed = _declared(_docx(), "data.docx")
    tables = _blocks_of(parsed, Extraction.TABLE)

    assert len(tables) == 1
    assert tables[0].text == "Region | Q3 | Q4\nR1C0 | R1C1 | R1C2\nR2C0 | R2C1 | R2C2"
    # The absence half, as this module's docstring demands of the PDF side: no cell may
    # survive in the prose around it, or the passage is indexed twice.
    assert "R1C0" not in _text_of(parsed, Extraction.TEXT)


def test_a_docx_table_declares_its_first_row_as_its_header() -> None:
    """python-docx's default template writes `w:tblLook w:firstRow="1"`, as Word does."""
    table = _blocks_of(_declared(_docx(), "data.docx"), Extraction.TABLE)[0]

    assert table.header == "Region | Q3 | Q4"
    assert table.text.split("\n")[0] == table.header  # the chunker re-reads line 0


def test_a_docx_tables_block_keeps_its_sections_locator() -> None:
    """R-88(6) refused a fourth `LocatorKind` for a PDF table; R-89 says the same here."""
    parsed = _declared(_docx(), "data.docx")
    table = _blocks_of(parsed, Extraction.TABLE)[0]

    assert table.locator.kind is LocatorKind.SECTION
    assert table.locator.section_path == ("Data",)
    # Prose, table and prose are one section yielding three blocks — not three sections.
    assert {block.locator.section_index for block in parsed.blocks} == {1}


def test_prose_before_and_after_a_docx_table_keep_body_order() -> None:
    parsed = _declared(_docx(), "data.docx")

    assert [block.extraction for block in parsed.blocks] == [
        Extraction.TEXT,
        Extraction.TABLE,
        Extraction.TEXT,
    ]
    assert parsed.blocks[0].text.endswith(PROSE_ABOVE)
    assert parsed.blocks[2].text == PROSE_BELOW
    assert [block.order for block in parsed.blocks] == [0, 1, 2]


def test_an_explicit_repeat_as_header_row_is_honoured_without_the_style_flag() -> None:
    """`w:tblHeader` is the strong signal — the author ticked "Repeat as header row"."""
    parsed = _declared(_docx(header_flag="", table_look=False), "data.docx")
    assert _blocks_of(parsed, Extraction.TABLE)[0].header == "Region | Q3 | Q4"


@pytest.mark.parametrize("value", ["0", "false"])
def test_a_header_row_flag_turned_off_is_not_a_header(value: str) -> None:
    """A toggle property that says off outranks the style flag that says on."""
    parsed = _declared(_docx(header_flag=value), "data.docx")
    assert _blocks_of(parsed, Extraction.TABLE)[0].header == ""


def test_a_table_with_neither_signal_declares_no_header_and_is_still_a_table_block() -> None:
    """Row atomicity does not depend on a header: an unlabelled grid is still a grid."""
    table = _blocks_of(_declared(_docx(table_look=False), "data.docx"), Extraction.TABLE)[0]

    assert table.header == ""
    assert table.extraction is Extraction.TABLE
    assert table.text.startswith("Region | Q3 | Q4")  # still line 0, just not *declared*


def test_a_one_row_docx_table_stays_prose() -> None:
    """One row is under the floor, and a header with nothing beneath it labels nothing."""
    parsed = _declared(_docx(rows=1), "data.docx")

    assert _blocks_of(parsed, Extraction.TABLE) == []
    assert "Region | Q3 | Q4" in _text_of(parsed, Extraction.TEXT)


def test_a_docx_table_whose_first_row_is_blank_declares_no_header() -> None:
    """The blank row is dropped, so the row that lands on line 0 is data, not a header."""
    table = _blocks_of(_declared(_docx(blank_first_row=True), "data.docx"), Extraction.TABLE)[0]

    assert table.header == ""
    assert table.text.startswith("R1C0")


def test_a_one_row_table_declares_no_header_even_below_a_lowered_floor() -> None:
    """`table_min_rows` is a `# TBD(§8.4)` knob, so "one row" is a reachable table.

    At the shipped floor of 2 a one-row table never becomes a table block at all, which makes
    the "a header needs a body row under it" guard look like dead code. It is not — it is one
    setting away, and a header repeated onto chunks of a block whose only line *is* that header
    would duplicate it into every citation.
    """
    docx = _blocks_of(_declared(_docx(rows=1), "one.docx", table_min_rows=1), Extraction.TABLE)
    assert docx[0].header == ""

    source = b"""# Data

| Region | Q3 |
| ------ | -- |
"""
    md = _blocks_of(_declared(source, "one.md", table_min_rows=1), Extraction.TABLE)
    assert md[0].header == ""


def test_a_docx_table_below_the_column_floor_stays_prose() -> None:
    """Below the floor the output is exactly what shipped before T-223 (R-89)."""
    parsed = _declared(_docx(cols=1), "data.docx")

    assert _blocks_of(parsed, Extraction.TABLE) == []
    assert "R1C0" in _text_of(parsed, Extraction.TEXT)


def test_a_split_docx_table_repeats_its_header_on_every_part() -> None:
    """`split_text` may cut the table before the chunker sees it — R-88(6)'s rule, here."""
    parsed = _declared(_docx(rows=40), "data.docx", max_block_chars=200)
    parts = _blocks_of(parsed, Extraction.TABLE)

    assert len(parts) > 1
    for part in parts:
        assert part.text.startswith("Region | Q3 | Q4")
        assert part.header == "Region | Q3 | Q4"
    # Part 0 already opens with the header — the parser put it there — so re-prefixing it
    # would print the column names twice. `startswith` cannot see that, which is why the
    # count is asserted separately (the same blind spot as the PDF twin above).
    assert parts[0].text.count("Region | Q3 | Q4") == 1


def test_a_docx_with_no_table_parses_exactly_as_before() -> None:
    """The blast-radius bound: T-223 changes nothing for a document with no table.

    The expected text is *composed* the way the pre-T-223 parser composed it — paragraphs
    joined by a blank line, normalised once — rather than read back out of a run of the new
    code, which would agree with itself whatever it did.
    """
    document = new_docx()
    document.add_heading("Data", level=1)
    for paragraph in ("First paragraph.", "Second paragraph.", "Third paragraph."):
        document.add_paragraph(paragraph)
    buffer = io.BytesIO()
    document.save(buffer)

    parsed = _declared(buffer.getvalue(), "flat.docx")
    expected = normalize(
        "\n\n".join(["Data", "First paragraph.", "Second paragraph.", "Third paragraph."])
    )

    assert [block.text for block in parsed.blocks] == [expected]
    assert parsed.blocks[0].extraction is Extraction.TEXT
    assert parsed.blocks[0].header == ""


def test_a_budget_breach_inside_a_docx_table_section_is_still_terminal() -> None:
    """Emitting per segment must not turn R-34's caps-reject rule into a partial parse."""
    with pytest.raises(DocumentTooComplexError):
        _declared(_docx(rows=40), "data.docx", max_extracted_chars=200)


# --- Markdown -----------------------------------------------------------------


def test_a_markdown_table_is_its_own_block_marked_table() -> None:
    parsed = _declared(MD_TABLE, "data.md")
    tables = _blocks_of(parsed, Extraction.TABLE)

    assert len(tables) == 1
    assert tables[0].text == "Region | Q3 | Q4\nEU | 12 | 15\nUS | 20 | 25"
    assert "EU | 12" not in _text_of(parsed, Extraction.TEXT)


def test_a_markdown_table_declares_its_thead_row_as_its_header() -> None:
    """GFM guarantees a header row and markdown-it reports it as `thead` — no guessing."""
    table = _blocks_of(_declared(MD_TABLE, "data.md"), Extraction.TABLE)[0]

    assert table.header == "Region | Q3 | Q4"
    assert table.text.split("\n")[0] == table.header


def test_a_markdown_table_keeps_its_sections_locator_and_line_range() -> None:
    """The recorded coarseness (R-89), asserted so it is a decision and not an accident."""
    parsed = _declared(MD_TABLE, "data.md")
    table = _blocks_of(parsed, Extraction.TABLE)[0]

    assert table.locator.kind is LocatorKind.SECTION
    assert {block.locator.section_index for block in parsed.blocks} == {1}
    assert len({(block.locator.line_start, block.locator.line_end) for block in parsed.blocks}) == 1


def test_prose_after_a_markdown_table_is_not_folded_back_into_the_prose_before_it() -> None:
    """The anti-fusion claim: `add` may never merge across a table segment."""
    parsed = _declared(MD_TABLE, "data.md")

    assert [block.extraction for block in parsed.blocks] == [
        Extraction.TEXT,
        Extraction.TABLE,
        Extraction.TEXT,
    ]
    assert parsed.blocks[0].text.endswith(PROSE_ABOVE)
    assert parsed.blocks[2].text == PROSE_BELOW


def test_two_adjacent_markdown_tables_stay_two_blocks() -> None:
    """Merging them would file the second table's rows under the first one's header."""
    source = b"""# Data

| Region | Q3 |
| ------ | -- |
| EU     | 12 |

| Country | Q4 |
| ------- | -- |
| US      | 25 |
"""
    tables = _blocks_of(_declared(source, "two.md"), Extraction.TABLE)

    assert len(tables) == 2
    assert tables[0].header == "Region | Q3"
    assert tables[1].header == "Country | Q4"


def test_a_markdown_table_with_no_header_text_declares_none() -> None:
    source = b"""# Data

|    |    |
| -- | -- |
| EU | 12 |
| US | 20 |
"""
    table = _blocks_of(_declared(source, "bare.md"), Extraction.TABLE)[0]

    assert table.header == ""
    assert table.text.startswith("EU | 12")


def test_a_markdown_cells_inner_whitespace_is_collapsed_like_every_other_format() -> None:
    """T-223 repointed `markdown.py` at `render_row`, which is half of the version bump.

    Before it, Markdown was the one format that kept an inner whitespace run — so the same
    table exported as `.md` and as `.docx` hashed differently for no reason a reader could see.
    """
    source = b"""# Data

| Region     | Q3 |
| ---------- | -- |
| EU   north | 12 |
| US   south | 20 |
"""
    table = _blocks_of(_declared(source, "ws.md"), Extraction.TABLE)[0]

    assert "EU north | 12" in table.text
    assert "EU   north" not in table.text


def test_the_markdown_delimiter_row_never_reaches_the_text() -> None:
    """**A contract with markdown-it, not a test of our code**, and worth saying so.

    Its `table` rule consumes the `|---|` line and emits no token for it, so this passes
    against anything we could write and would survive deleting all of it — it is not coverage
    (T-219's lesson about a test that cannot fail for the reason it names). It is kept because
    a library bump that started emitting that row would otherwise put `--- | ---` into an
    embedding, silently, and nothing else in the suite would notice.
    """
    assert "---" not in _declared(MD_TABLE, "data.md").text


def test_a_markdown_table_below_the_row_floor_stays_prose() -> None:
    source = b"""# Data

| Region | Q3 |
| ------ | -- |
"""
    parsed = _declared(source, "thin.md")

    assert _blocks_of(parsed, Extraction.TABLE) == []
    assert "Region | Q3" in _text_of(parsed, Extraction.TEXT)
